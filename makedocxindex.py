from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, RGBColor, Mm

def create_toc_docx(toc_entries, casedetails, output_file_path, confidential=False, date_setting=True, index_font_setting=None):
    # Create a new Word document
    doc = Document()

    style = doc.styles['Normal']
    #parse font setting same as in bundle.py:
    # if font = 
    # style.font.name = 'Times New Roman'
    if index_font_setting == "sans":
        style.font.name = "Arial"
    elif index_font_setting == "serif":
        style.font.name = "Times New Roman"
    elif index_font_setting == "mono":
        style.font.name = "Courier New"
    else:
        style.font.name = "Times New Roman"

    # Set up case details
    claimno_hdr = casedetails[1] if casedetails[1] else ""
    casename = casedetails[2] if casedetails[2] else ""
    bundle_name = casedetails[0].upper() if casedetails[0] else ""

    # Add the Claim Number (if exists)
    if claimno_hdr:
        para = doc.add_paragraph(claimno_hdr)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add Case Name (if exists)
    if casename:
        para = doc.add_paragraph(casename)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(14)

    # Add the Bundle Name
    if bundle_name:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(bundle_name)
        run.bold = True
        run.font.size = Pt(16)
        if confidential:
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.text = f"CONFIDENTIAL\n{bundle_name}"

# Table of Contents header row
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'  
    table.autofit = False
    table.allow_autofit = False
    
    # Set preferred widths in millimeters
    width_distribution = [Mm(14), Mm(95), Mm(35), Mm(17)]
        # ('12mm', 12),
        # ('100mm', 100),
        # ('35mm', 35),
        # ('12mm', 12)
    #]

    for row in table.rows:
        for idx, width in enumerate(width_distribution):
            table.columns[idx].width=(width)

# # Set fixed layout and overall table width
#     tbl = table._element
#     tblPr = tbl.get_or_add_tblPr()
#     tblLayout = OxmlElement('w:tblLayout')
#     tblLayout.set(qn('w:type'), 'fixed')
#     tblPr.append(tblLayout)

#     tcW = OxmlElement('w:tcW')
#     tcW.set(qn('w:w'), str(int(159 * 50800 / 25.4)))  # Total width in twips
#     tcW.set(qn('w:type'), 'dxa')
#     tblPr.append(tcW)

#     # Set column widths
#     for row in table.rows:
#         for idx, cell in enumerate(row.cells):
#             tcPr = cell._element.get_or_add_tcPr()
#             tcW = OxmlElement('w:tcW')
#             tcW.set(qn('w:w'), str(int(width_distribution[idx][1] * 50800 / 25.4)))  # Convert mm to twips
#             tcW.set(qn('w:type'), 'dxa')
#             tcPr.append(tcW)


    header_cells = table.rows[0].cells
    header_cells[0].text = "Tab"
    header_cells[1].text = "Title"
    header_cells[2].text = "Date" if date_setting else ""
    header_cells[3].text = "Page"
    for cell in header_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    # Add entries to the Table of Contents
    for entry in toc_entries:
        row = table.add_row().cells
        if "SECTION_BREAK" in entry[0]:
            # Handle section breaks
            para = doc.add_paragraph(entry[1])
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if para.runs:
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(12)
            row[0].merge(row[-1])  # Merge all cells for a section header
        else:
            # Add a regular TOC entry
            row[0].text = str(entry[0])  # Tab
            row[1].text = entry[1]      # Title
            row[2].text = entry[2] if date_setting else ""  # Date
            row[3].text = str(entry[3])  # Page

    # Save the document
    doc.save(output_file_path)

# main
if __name__ == "__main__":
    # Sample data
    toc_entries = [
        ("001.", "First Doc", "2021-01-01", 1),
        ("002", "Second Doc", "2021-01-02", 5),
        ("003", "Third Document", "2021-01-03", 10),
        #("SECTION_BREAK", "Section Break Test", "", ""),
        ("004", "Document Number Four", "2021-01-04", 15),
        ("005", "The fifth document in this series", "2021-01-05", 20),
    ]
    casedetails = ["Bundle Name", "Claim Number", "Case Name"]
    output_file_path = "TOC.docx"

    # Create the TOC document
    create_toc_docx(toc_entries, casedetails, output_file_path, confidential=True, date_setting=False)
    print(f"Table of Contents saved to '{output_file_path}'")